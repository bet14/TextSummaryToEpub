import requests
import json
import os
import time
from datetime import datetime
from docx import Document
from ebooklib import epub

# Đọc API key từ file
api_key_path = r"D:\Code\API KEY\Gemini API Key.txt"
try:
    with open(api_key_path, "r", encoding="utf-8") as key_file:
        api_key = key_file.read().strip()
except FileNotFoundError:
    print(f"Không tìm thấy file API key tại {api_key_path}")
    exit(1)

# Thư mục chứa file txt cần đọc và thư mục output
input_dir = r"D:\2025 Archive\Raw Articles txt"
today_str = datetime.now().strftime("%Y-%m-%d")
output_dir = os.path.join(r"D:\2025 Archive\Summaries", f"{today_str} - Summaries")
output_docx_path = os.path.join(r"D:\2025 Archive\Summaries", f"{today_str}_summaries.docx")
output_epub_path = os.path.splitext(output_docx_path)[0] + '.epub'

# Hàm khởi tạo URL API cho từng model
def get_api_url(key, model_version):
    if model_version == "2.5":
        return f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={key}"
    else:
        raise ValueError(f"Phiên bản model {model_version} không hợp lệ")

# Cấu hình API và luân phiên model
model_versions = ["2.5", "2.5"]
current_model_index = 0
url = get_api_url(api_key, model_versions[current_model_index])
headers = {"Content-Type": "application/json"}
summary_count = 0
max_summaries_per_chat = 1

# Hàm lấy tất cả file .txt trong thư mục và subfolder
def get_all_txt_files(root_dir):
    txt_files = []
    for dirpath, dirnames, filenames in os.walk(root_dir):
        for filename in filenames:
            if filename.lower().endswith(".txt"):
                full_path = os.path.join(dirpath, filename)
                relative_dir = os.path.relpath(dirpath, root_dir)
                txt_files.append((full_path, relative_dir))
    return txt_files

# Hàm tạo tóm tắt từ nội dung file
def generate_summary(file_path, relative_dir):
    global current_model_index, summary_count, url

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text_content = f.read()
    except Exception as e:
        print(f"Không thể đọc file {os.path.basename(file_path)}: {str(e)}")
        return None

    word_count_original = len(text_content.split())
    if word_count_original < 50:
        print(f"Bỏ qua file {os.path.basename(file_path)} vì chỉ có {word_count_original} chữ (dưới 50 chữ).")
        return None

    # Tạo prompt yêu cầu tóm tắt
    prompt_text = (
        "Yêu cầu tóm tắt chi tiết bằng tiếng Việt, trên 5000 từ, không thêm lời giới thiệu hay nhận xét, "
        "không giải thích đã làm gì, không đề cập số từ đã tóm tắt, chỉ cung cấp nội dung tóm tắt. "
        "BẮT BUỘC phải trên 2000 từ, trình bày tóm tắt theo định dạng: đầu tiên là dòng Tên bài bằng tiếng Việt, "
        "sau đó là các đoạn văn, không dùng gạch đầu dòng hay liệt kê. "
        f"Nội dung bài viết:\n{text_content}\n\n"
    )

    data = {
        "contents": [{"parts": [{"text": prompt_text}]}]
    }

    # Gửi yêu cầu API
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
    except requests.RequestException as e:
        print(f"Lỗi khi xử lý file {os.path.basename(file_path)}: {str(e)}")
        return None

    result = response.json()
    summary_dict = result.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0]
    summary = summary_dict.get("text", "") if isinstance(summary_dict, dict) else summary_dict
    summary_clean = summary.replace("*", "").replace("#", "").strip()

    word_count_summary = len(summary_clean.split())
    print(f"File: {os.path.basename(file_path)}")
    print(f"Bài gốc: {word_count_original} chữ")
    print(f"Bài tóm tắt: {word_count_summary} chữ")
    print("Nội dung tóm tắt (100 chữ đầu):")
    print(' '.join(summary_clean.split()[:100]))

    # Tạo thư mục output tương ứng
    output_subdir = os.path.join(output_dir, relative_dir)
    os.makedirs(output_subdir, exist_ok=True)

    # Lưu file tóm tắt
    output_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_summary.txt"
    output_file_path = os.path.join(output_subdir, output_filename)
    try:
        with open(output_file_path, "w", encoding="utf-8") as f_out:
            f_out.write(summary_clean)
    except Exception as e:
        print(f"Không thể lưu tóm tắt cho {os.path.basename(file_path)}: {str(e)}")
        return None

    summary_count += 1
    if summary_count >= max_summaries_per_chat:
        current_model_index = (current_model_index + 1) % len(model_versions)
        url = get_api_url(api_key, model_versions[current_model_index])
        summary_count = 0

    return output_file_path

# Hàm tạo file DOCX tổng hợp
def create_docx_summary(summary_files):
    doc = Document()
    remove_sentence = "Dưới đây là bản tóm tắt chi tiết nội dung của đoạn văn bằng tiếng Việt"

    for file_path in summary_files:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except Exception as e:
            print(f"Không thể đọc file tóm tắt {os.path.basename(file_path)}: {str(e)}")
            continue

        content_lines = content.splitlines()
        filtered_lines = [line for line in content_lines if remove_sentence.lower() not in line.lower()]
        while filtered_lines and not filtered_lines[0].strip():
            filtered_lines.pop(0)

        title_line_index = -1
        for i, line in enumerate(filtered_lines):
            if line.strip().lower().startswith("tên bài"):
                title_line_index = i
                break

        if title_line_index >= 0:
            for i, line in enumerate(filtered_lines):
                line = line.strip()
                if not line:
                    continue
                if i == title_line_index:
                    after_prefix = line[len("Tên bài"):].lstrip()
                    if after_prefix.startswith(":"):
                        after_prefix = after_prefix[1:].lstrip()
                    doc.add_heading(after_prefix, level=1)
                else:
                    doc.add_paragraph(line)
        else:
            contain_title_index = -1
            for i, line in enumerate(filtered_lines):
                if "tên bài" in line.lower():
                    contain_title_index = i
                    break

            if contain_title_index >= 0:
                heading_text = filtered_lines[contain_title_index].strip()
                doc.add_heading(heading_text, level=1)
                for i, line in enumerate(filtered_lines):
                    if i == contain_title_index:
                        continue
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)
            else:
                first_text_line_index = -1
                for i, line in enumerate(filtered_lines):
                    if line.strip():
                        first_text_line_index = i
                        break

                if first_text_line_index >= 0:
                    heading_text = filtered_lines[first_text_line_index].strip()
                    doc.add_heading(heading_text, level=1)
                    for i, line in enumerate(filtered_lines):
                        if i == first_text_line_index:
                            continue
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line)

        doc.add_paragraph()

    try:
        doc.save(output_docx_path)
        print(f"Đã tạo file tổng hợp DOCX: {output_docx_path}")
    except Exception as e:
        print(f"Không thể lưu file DOCX: {str(e)}")
        return False
    return True

# Hàm chuyển DOCX sang EPUB
def docx_to_epub_with_bookmarks(docx_path, epub_path):
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Không thể đọc file DOCX {docx_path}: {str(e)}")
        return

    book = epub.EpubBook()
    book.set_identifier('Tóm tắt bằng AI')
    book.set_language('vi')

    chapters = []
    toc = []
    chapter_count = 0
    current_chapter_content = []
    current_chapter_title = None

    def add_chapter(title, content):
        nonlocal chapter_count
        chapter_count += 1
        chapter = epub.EpubHtml(title=title, file_name=f'chap_{chapter_count}.xhtml', lang='vi')
        html_content = f'<h1>{title}</h1>'
        for para in content:
            text = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            html_content += f'<p>{text}</p>'
        chapter.content = html_content.encode('utf-8')
        book.add_item(chapter)
        return chapter

    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()
        if not text:
            continue

        if style == 'Heading 1':
            if current_chapter_title is not None:
                chapter = add_chapter(current_chapter_title, current_chapter_content)
                chapters.append(chapter)
                toc.append(epub.Link(chapter.file_name, current_chapter_title, f'chap_{chapter_count}'))
            current_chapter_title = text
            current_chapter_content = []
        else:
            if current_chapter_title is not None:
                current_chapter_content.append(text)

    if current_chapter_title is not None:
        chapter = add_chapter(current_chapter_title, current_chapter_content)
        chapters.append(chapter)
        toc.append(epub.Link(chapter.file_name, current_chapter_title, f'chap_{chapter_count}'))

    book.toc = toc
    book.spine = ['nav'] + chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    try:
        epub.write_epub(epub_path, book, {})
        print(f"Đã tạo file EPUB: {epub_path}")
    except Exception as e:
        print(f"Không thể tạo file EPUB: {str(e)}")

# Main execution
def main():
    file_list = get_all_txt_files(input_dir)
    summary_files = []

    for file_path, relative_dir in file_list:
        summary_file = generate_summary(file_path, relative_dir)
        if summary_file:
            summary_files.append(summary_file)
        time.sleep(5)  # Chờ trước khi xử lý file tiếp theo

    if summary_files:
        if create_docx_summary(summary_files):
            docx_to_epub_with_bookmarks(output_docx_path, output_epub_path)
    else:
        print("Không có file tóm tắt nào được tạo.")

if __name__ == "__main__":
    main()